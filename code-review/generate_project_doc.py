"""Generate a comprehensive project documentation .docx for the Go Code Review system.

Run from the `code-review/` directory:

    python generate_project_doc.py

Output: project_overview_go_code_reviewer.docx
"""

from datetime import datetime

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

OUTPUT_PATH = "project_overview_go_code_reviewer.docx"


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)


def main() -> None:
    doc = Document()

    # ---------- Title ----------
    p = doc.add_paragraph()
    r = p.add_run("Go Code Review System\nProject Overview & Architecture Reference")
    r.bold = True
    r.font.size = Pt(18)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    meta = doc.add_paragraph(
        f"Prepared on: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        "Scope: end-to-end documentation of every component in the code-review/ project."
    )
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ---------- 1. Executive Summary ----------
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "The Go Code Review System is an LLM-powered, RAG-augmented, multi-mode code "
        "reviewer for Go repositories. It detects violations of internal coding standards "
        "(security, error handling, concurrency, naming, observability, performance, "
        "documentation, testing) and emits structured findings in Markdown, JSON, or SARIF. "
        "It supports three review modes (hybrid, RAG-only, fine-tune-only) and two RAG "
        "strategies (simple single-shot retrieval and an agentic LangGraph multi-agent "
        "workflow)."
    )

    doc.add_heading("1.1 Capabilities at a glance", level=2)
    add_bullets(doc, [
        "Parses Go source via the official AST (with regex fallback) into function/method/struct chunks.",
        "Retrieves relevant rules from a Qdrant vector store built from internal standards.",
        "Drives inference through either a local fine-tuned model or a remote Ollama model.",
        "Optional agentic mode runs six specialist agents (security, architecture, performance, observability, database, concurrency) in parallel via LangGraph.",
        "Deduplicates and ranks findings by severity, then renders Markdown / JSON / SARIF.",
        "Ships with a 17-file seeded benchmark and an evaluator for precision, recall, F1.",
        "Exposes both a Click CLI and a FastAPI REST server.",
    ])

    # ---------- 2. Operating modes ----------
    doc.add_heading("2. Operating Modes", level=1)
    add_table(doc,
        ["Flag", "Mode", "Description"],
        [
            ["--mode hybrid", "Hybrid", "RAG retrieval + fine-tuned model inference"],
            ["--mode rag-only", "RAG-only", "Retrieval + base/Ollama model (no fine-tuning needed)"],
            ["--mode fine-tune-only", "Fine-tune only", "Fine-tuned model only, no retrieval"],
            ["--rag-mode simple", "Simple RAG", "Single-shot top-k retrieval per chunk"],
            ["--rag-mode agentic", "Agentic RAG", "LangGraph multi-agent (planner → 6 agents → aggregator)"],
        ])

    # ---------- 3. Repository layout ----------
    doc.add_heading("3. Repository Layout", level=1)
    add_table(doc,
        ["Path", "Purpose"],
        [
            ["pipeline/", "Core orchestrator, Go AST chunking, dedup, report formatting"],
            ["rag/", "Vector store builder, retriever, agent tools, LangGraph orchestrator"],
            ["training/", "QLoRA fine-tuning script and evaluation harnesses"],
            ["dataset/", "Dataset construction from linter output and manual examples"],
            ["cli/", "Click-based CLI front-end"],
            ["serving/", "FastAPI REST server + Dockerfile"],
            ["standards/", "Markdown standards docs and rules.json"],
            ["benchmarks/", "Seeded Go repo (17 files + 1 clean) and ground_truth.json"],
            ["results/", "Pipeline output and evaluation reports"],
        ])

    # ---------- 4. End-to-end workflow ----------
    doc.add_heading("4. End-to-End Workflow", level=1)
    doc.add_paragraph(
        "The diagram below summarizes how a request flows through the system. "
        "Steps in brackets are skipped depending on the chosen mode."
    )
    add_code(doc, (
        "CLI / API request\n"
        "   |\n"
        "   v\n"
        "GoReviewPipeline.load()\n"
        "   |- load model (local fine-tuned OR Ollama)\n"
        "   |- load RAG retriever (Qdrant + nomic-embed-text-v1.5)\n"
        "   '- [agentic] load MultiAgentReviewer (LangGraph)\n"
        "   |\n"
        "   v\n"
        "extract_go_chunks(repo)        <- pipeline/go_parser.py (Go AST helper or regex)\n"
        "   |\n"
        "   v\n"
        "for each GoChunk:\n"
        "   |- simple   : retriever.retrieve(code) -> build_prompt -> _generate -> parse_findings\n"
        "   |- agentic  : multi_agent.review_code(code) -> planner -> agents (parallel) -> aggregator\n"
        "   '- ft-only  : build_prompt without rules -> _generate\n"
        "   |\n"
        "   v\n"
        "deduplicate_findings -> rank_findings -> compute_summary\n"
        "   |\n"
        "   v\n"
        "RepositoryReport\n"
        "   |\n"
        "   v\n"
        "generate_markdown_report / generate_json_report / generate_sarif_report\n"
    ))

    # ---------- 5. Modules ----------
    doc.add_heading("5. Module-by-Module Reference", level=1)

    # 5.1 review_pipeline
    doc.add_heading("5.1 pipeline/review_pipeline.py — orchestrator", level=2)
    doc.add_paragraph(
        "Core orchestration class GoReviewPipeline plus dataclasses ReviewConfig, "
        "ReviewFinding, ReviewResult, RepositoryReport. Loads model and retriever, "
        "iterates Go chunks, builds prompts, calls the LLM, parses output, and "
        "produces the final report."
    )
    add_table(doc,
        ["Method", "Purpose"],
        [
            ["load()", "Init model, tokenizer, retriever, optional MultiAgentReviewer"],
            ["close()", "Free GPU memory, close Qdrant client"],
            ["review_repository(repo_path)", "Main API: review entire repo"],
            ["review_chunk(code, file, name)", "Review a single chunk"],
            ["_load_local_model()", "Load fine-tuned model + LoRA adapter or merged model"],
            ["_load_ollama()", "Verify Ollama reachable + model pulled"],
            ["_build_prompt(code, rules)", "Compose user prompt from code + retrieved rules"],
            ["_generate(system, user)", "Dispatch to local model or Ollama API"],
            ["_load_static_rules()", "Fallback when RAG unavailable: load rules.json"],
        ])

    doc.add_paragraph("Key configuration knobs (ReviewConfig):")
    add_table(doc,
        ["Field", "Default", "Purpose"],
        [
            ["mode", "hybrid", "hybrid | rag-only | fine-tune-only"],
            ["rag_mode", "simple", "simple | agentic"],
            ["ollama_model", "None", "If set, use Ollama remote inference"],
            ["ollama_url", "http://localhost:11434", "Ollama base URL"],
            ["max_new_tokens", "2048", "LLM output length cap"],
            ["temperature", "0", "0 = greedy / deterministic"],
            ["top_p", "1.0", "nucleus sampling"],
            ["seed", "42", "fixed seed for reproducibility"],
            ["top_k", "5", "rules retrieved per chunk"],
            ["max_chunk_tokens", "3000", "input length cap"],
            ["use_quantization", "True", "4-bit QLoRA loading"],
            ["debug", "False", "print raw LLM output per chunk"],
        ])

    # 5.2 go_parser
    doc.add_heading("5.2 pipeline/go_parser.py — Go AST chunking", level=2)
    doc.add_paragraph(
        "Splits a Go repo into reviewable units (functions, methods, structs, interfaces). "
        "Primary path: compile and run a small Go helper (ast_helper.go) that uses go/ast "
        "to emit JSON with exact line ranges. Fallback: regex matching when the Go "
        "toolchain is not installed."
    )
    doc.add_paragraph(
        "GoChunk fields: file_path, chunk_type, name, package, start_line, end_line, "
        "code, imports, receiver, doc_comment."
    )
    doc.add_paragraph(
        "Note: the helper is compiled into tempfile.gettempdir()/go_code_reviewer/. "
        "On shared hosts, set TMPDIR to a user-writable directory before running."
    )

    # 5.3 deduplication
    doc.add_heading("5.3 pipeline/deduplication.py — parsing, dedup, ranking", level=2)
    add_table(doc,
        ["Function", "Purpose"],
        [
            ["parse_findings(text)", "Regex-extract structured findings from LLM output"],
            ["deduplicate_findings(findings)", "Remove near-duplicates (same rule + file + line)"],
            ["rank_findings(findings)", "Sort by severity then line number"],
            ["compute_summary(findings)", "Severity counts dict"],
            ["group_findings_by_file(findings)", "For per-file report sections"],
            ["group_findings_by_category(findings)", "For per-category report sections"],
        ])

    # 5.4 report_generator
    doc.add_heading("5.4 pipeline/report_generator.py — output formats", level=2)
    add_table(doc,
        ["Function", "Format"],
        [
            ["generate_markdown_report(report)", "Markdown (overview, severity table, by-file, by-category)"],
            ["generate_json_report(report)", "JSON with raw output and config"],
            ["generate_sarif_report(report)", "SARIF 2.1.0 for GitHub / Azure DevOps"],
        ])

    # 5.5 retriever
    doc.add_heading("5.5 rag/retriever.py — vector retrieval", level=2)
    doc.add_paragraph(
        "GoStandardsRetriever wraps a local Qdrant collection (default: go_coding_standards) "
        "and HuggingFace embeddings (nomic-ai/nomic-embed-text-v1.5, 768-dim). Methods: "
        "retrieve, retrieve_with_scores, format_rules_for_prompt, close. Includes a "
        "compatibility shim for newer Qdrant clients that only expose query_points()."
    )

    # 5.6 build_vector_store
    doc.add_heading("5.6 rag/build_vector_store.py — index standards", level=2)
    add_bullets(doc, [
        "Loads markdown standards via DirectoryLoader.",
        "Loads rules.json — each rule becomes a Document with metadata.",
        "Splits with MarkdownTextSplitter (500 tokens, 50 overlap).",
        "Embeds with nomic-embed-text-v1.5.",
        "Ingests into Qdrant (cosine, 768-dim) and persists to disk.",
    ])

    # 5.7 agent_tools
    doc.add_heading("5.7 rag/agent_tools.py — deterministic tool functions", level=2)
    doc.add_paragraph(
        "Plain Python functions used by the agentic reviewer. Tools execute in Python "
        "(not via model tool-calling); their outputs are materialized into the prompt "
        "before the LLM is invoked."
    )
    add_table(doc,
        ["Tool", "Purpose"],
        [
            ["tool_search_code", "Regex search across .go files"],
            ["tool_read_file", "Read a file (<= 200 KB)"],
            ["tool_list_directory", "List directory contents"],
            ["tool_read_go_mod", "Read go.mod"],
            ["tool_parse_ast", "Lightweight AST summary (imports, funcs, types, goroutine/channel counts)"],
            ["tool_query_dependency_graph", "Build file -> imports map"],
            ["tool_run_golangci_lint", "Run golangci-lint if installed"],
            ["tool_run_go_vet", "Run go vet if Go toolchain present"],
            ["tool_query_schema", "Collect DDL from .sql files"],
            ["tool_explain_query", "Static heuristics over SQL (SELECT *, missing WHERE, N+1)"],
            ["tool_build_toolbox", "Bundle all tools into a dict"],
        ])
    doc.add_paragraph(
        "All tools degrade gracefully — missing binaries return [tool-unavailable]."
    )

    # 5.8 langgraph_agents
    doc.add_heading("5.8 rag/langgraph_agents.py — multi-agent orchestration", level=2)
    doc.add_paragraph(
        "Six specialist agents coordinated via a LangGraph StateGraph. Planner picks "
        "2 to 5 relevant agents per chunk; agents execute in parallel; aggregator "
        "deduplicates and ranks the union of their findings."
    )
    add_table(doc,
        ["Agent", "Persona", "Focus", "Rule IDs"],
        [
            ["security", "AppSec Engineer", "SQLi, auth, secrets, input validation", "SEC-001/2/3"],
            ["architecture", "Principal Architect", "service boundaries, DI, contracts", "NAM-001/2, DOC-001"],
            ["performance", "Performance Engineer", "N+1, allocations, leaks", "PERF-001"],
            ["observability", "SRE / Platform", "logging, traces, metrics, health", "LOG-001/2"],
            ["database", "DB Reliability", "tx safety, batches, indexes", "SEC-002"],
            ["concurrency", "Systems Engineer", "goroutine lifecycle, channels, mutexes", "CONC-001/2, CTX-001/2"],
        ])
    add_code(doc,
        "START -> planner -> [security | architecture | performance |\n"
        "                     observability | database | concurrency]   (parallel)\n"
        "                  -> aggregator -> END"
    )
    doc.add_paragraph(
        "Each agent makes exactly one LLM call per chunk; tool outputs are concatenated "
        "into the prompt up front."
    )

    # 5.9 fine_tune
    doc.add_heading("5.9 training/fine_tune_go_reviewer.py — QLoRA fine-tuning", level=2)
    add_bullets(doc, [
        "Base model: deepseek-ai/deepseek-coder-7b-instruct-v1.5 (Qwen2.5-Coder also supported).",
        "Quantization: 4-bit NF4 with double quantization, bf16 compute.",
        "LoRA: r=64, alpha=128, dropout=0.05; targets q/k/v/o_proj, gate/up/down_proj.",
        "Training: 3 epochs, per-device batch=2, grad_accum=4, lr=2e-4 cosine, 5% warmup.",
        "max_seq_length=2048; eval/save per epoch.",
        "Outputs: LoRA adapter at output_dir/, merged model at final_model_dir/.",
    ])

    # 5.10 evaluate_seeded
    doc.add_heading("5.10 training/evaluate_seeded.py — seeded benchmark eval", level=2)
    doc.add_paragraph(
        "Compares pipeline output to benchmarks/ground_truth.json. Per-file outcome: "
        "PASS / MISS / NOISE / PARTIAL. Aggregate metrics: precision, recall, F1, "
        "false-positive rate."
    )

    # 5.11 evaluate
    doc.add_heading("5.11 training/evaluate.py — held-out evaluation", level=2)
    doc.add_paragraph(
        "Runs the fine-tuned model over a JSONL test set, computing per-category and "
        "aggregate precision/recall/F1."
    )

    # 5.12 build_dataset
    doc.add_heading("5.12 dataset/build_dataset.py — dataset construction", level=2)
    doc.add_paragraph(
        "Converts linter outputs (golangci-lint, gosec) and manual examples into JSONL "
        "chat training data. Maps linter codes to internal rule IDs (e.g. errcheck -> "
        "EH-002, gosec G101 -> SEC-001) and extracts ~15 lines of context per violation."
    )

    # 5.13 cli
    doc.add_heading("5.13 cli/review.py — Click CLI front-end", level=2)
    doc.add_paragraph(
        "Subcommands: repo, file, snippet, build-rag, serve. Common flags: --model, "
        "--base-model, --mode, --rag-db, --format, --output, --ollama, --no-quant, --quiet."
    )

    # 5.14 api
    doc.add_heading("5.14 serving/api.py — FastAPI REST server", level=2)
    add_table(doc,
        ["Endpoint", "Method", "Purpose"],
        [
            ["/health", "GET", "Liveness/readiness"],
            ["/review/code", "POST", "Review a snippet"],
            ["/review/file", "POST", "Review uploaded .go file"],
            ["/review/repo", "POST", "Review a repo path on the server"],
            ["/rules", "GET", "Dump all standards rules"],
        ])
    doc.add_paragraph(
        "Pipeline is loaded once on app startup. Configuration via env vars: "
        "MODEL_PATH, BASE_MODEL, RAG_DB_PATH, REVIEW_MODE, OLLAMA_MODEL, OLLAMA_URL, "
        "USE_QUANTIZATION."
    )

    # ---------- 6. Standards ----------
    doc.add_heading("6. Standards & Rule Catalog", level=1)
    doc.add_paragraph(
        "Source: standards/ directory. Markdown docs are human-readable; rules.json "
        "is the structured form (rule_id, category, severity, title, description, "
        "violation_example, correct_example, reference, auto_fixable)."
    )

    doc.add_heading("6.1 Security", level=2)
    add_table(doc, ["Rule", "Severity", "Description"], [
        ["SEC-001", "CRITICAL", "No hardcoded secrets (creds, API keys, tokens)"],
        ["SEC-002", "CRITICAL", "Parameterized SQL queries, no string concat"],
        ["SEC-003", "CRITICAL", "Validate / sanitize all external input"],
        ["SEC-004", "HIGH", "Use TLS for network I/O"],
        ["SEC-005", "HIGH", "Set timeouts on HTTP servers/clients"],
    ])

    doc.add_heading("6.2 Error handling", level=2)
    add_table(doc, ["Rule", "Severity", "Description"], [
        ["EH-001", "HIGH", "Wrap errors with fmt.Errorf(\"...: %w\", err)"],
        ["EH-002", "HIGH", "Never ignore errors"],
        ["EH-003", "HIGH", "No panic in library/service code"],
        ["EH-004", "MEDIUM", "Sentinel errors via errors.Is"],
    ])

    doc.add_heading("6.3 Concurrency", level=2)
    add_table(doc, ["Rule", "Severity", "Description"], [
        ["CONC-001", "HIGH", "Goroutines must have a termination path"],
        ["CONC-002", "MEDIUM", "Use directional channels (chan<-, <-chan)"],
        ["CONC-003", "HIGH", "Protect shared state with sync.Mutex"],
        ["CONC-004", "MEDIUM", "Use sync.WaitGroup for fan-out"],
    ])

    doc.add_heading("6.4 Naming", level=2)
    add_table(doc, ["Rule", "Severity", "Description"], [
        ["NAM-001", "LOW", "MixedCaps, no underscores"],
        ["NAM-002", "LOW", "Acronyms all-caps (HTTPClient, userID)"],
        ["NAM-003", "LOW", "Short receiver names"],
        ["NAM-004", "LOW", "Single-method interfaces end with -er"],
        ["NAM-005", "LOW", "Lowercase single-word package names"],
    ])

    doc.add_heading("6.5 Context, logging, perf, docs, tests", level=2)
    add_table(doc, ["Rule", "Severity", "Description"], [
        ["CTX-001", "HIGH", "context.Context is the first parameter"],
        ["CTX-002", "MEDIUM", "Never store context in a struct"],
        ["LOG-001", "MEDIUM", "No fmt.Println; use structured logger"],
        ["LOG-002", "MEDIUM", "No fmt.Printf; use structured logger"],
        ["PERF-001", "MEDIUM", "Pre-allocate slices when size is known"],
        ["DOC-001", "LOW", "Exported funcs require doc comments"],
        ["TEST-001", "MEDIUM", "Tests should be table-driven"],
    ])

    # ---------- 7. Benchmark ----------
    doc.add_heading("7. Seeded Benchmark", level=1)
    doc.add_paragraph(
        "benchmarks/seeded_repo/ contains 18 Go files (17 with a single seeded "
        "violation each + 1 clean negative test). benchmarks/ground_truth.json maps "
        "every file to its expected rule_id and severity."
    )
    add_table(doc, ["File", "Expected rule"], [
        ["sec001_hardcoded.go", "SEC-001 / CRITICAL"],
        ["sec002_sqli.go", "SEC-002 / CRITICAL"],
        ["sec003_jwt_secret.go", "SEC-003 / CRITICAL"],
        ["eh001_bare_return.go", "EH-001 / HIGH"],
        ["eh002_ignored_error.go", "EH-002 / HIGH"],
        ["eh003_panic.go", "EH-003 / HIGH"],
        ["ctx001_ctx_not_first.go", "CTX-001 / HIGH"],
        ["ctx002_ctx_in_struct.go", "CTX-002 / MEDIUM"],
        ["log001_println.go", "LOG-001 / MEDIUM"],
        ["log002_printf.go", "LOG-002 / MEDIUM"],
        ["nam001_underscore.go", "NAM-001 / LOW"],
        ["nam002_acronym.go", "NAM-002 / LOW"],
        ["conc001_goroutine.go", "CONC-001 / HIGH"],
        ["conc002_bidirectional.go", "CONC-002 / MEDIUM"],
        ["test001_no_table.go", "TEST-001 / MEDIUM"],
        ["perf001_slice_alloc.go", "PERF-001 / MEDIUM"],
        ["doc001_no_comment.go", "DOC-001 / LOW"],
        ["clean_code.go", "(none — negative test)"],
    ])

    # ---------- 8. Data formats ----------
    doc.add_heading("8. Data Formats", level=1)

    doc.add_heading("8.1 Training example (JSONL line)", level=2)
    add_code(doc, (
        '{\n'
        '  "messages": [\n'
        '    {"role": "system",    "content": "You are an expert Go code reviewer..."},\n'
        '    {"role": "user",      "content": "## Coding Standards...\\n\\n## Go Code\\n```go\\n...\\n```"},\n'
        '    {"role": "assistant", "content": "### VIOLATION [EH-001] HIGH — Bare return\\n..."}\n'
        '  ]\n'
        '}'
    ))

    doc.add_heading("8.2 Repository report (JSON)", level=2)
    add_code(doc, (
        '{\n'
        '  "repo_path": "/path/to/repo",\n'
        '  "total_files": 12,\n'
        '  "total_chunks": 48,\n'
        '  "total_findings": 23,\n'
        '  "findings": [\n'
        '    {\n'
        '      "rule_id": "SEC-001", "severity": "CRITICAL", "category": "security",\n'
        '      "title": "Hardcoded credentials", "file": "config.go",\n'
        '      "line_start": 15, "line_end": 15, "function": "loadConfig",\n'
        '      "description": "API key hardcoded in source",\n'
        '      "current_code":  "const apiKey = \\"sk-abc123...\\"",\n'
        '      "suggested_fix": "apiKey := os.Getenv(\\"API_KEY\\")",\n'
        '      "effort": "easy", "auto_fixable": true\n'
        '    }\n'
        '  ],\n'
        '  "summary": {"critical": 1, "high": 5, "medium": 8, "low": 9, "info": 0},\n'
        '  "elapsed_seconds": 45.3,\n'
        '  "config": {"mode": "hybrid", "rag_mode": "simple", "top_k": 5}\n'
        '}'
    ))

    # ---------- 9. Typical commands ----------
    doc.add_heading("9. Typical Commands", level=1)
    doc.add_paragraph("Run these from the code-review/ directory.")

    doc.add_heading("9.1 Build the vector store (one-time)", level=2)
    add_code(doc, "python rag/build_vector_store.py --rebuild")

    doc.add_heading("9.2 Simple RAG review", level=2)
    add_code(doc,
        "python pipeline/review_pipeline.py \\\n"
        "    --repo ./benchmarks/seeded_repo \\\n"
        "    --output ./results/seeded_review_simple.json \\\n"
        "    --mode rag-only \\\n"
        "    --ollama-model qwen2.5-coder:7b"
    )

    doc.add_heading("9.3 Agentic RAG review", level=2)
    add_code(doc,
        "python pipeline/review_pipeline.py \\\n"
        "    --repo ./benchmarks/seeded_repo \\\n"
        "    --output ./results/seeded_review_agentic.json \\\n"
        "    --mode rag-only --rag-mode agentic \\\n"
        "    --ollama-model qwen2.5-coder:7b"
    )

    doc.add_heading("9.4 Evaluate against ground truth", level=2)
    add_code(doc,
        "python training/evaluate_seeded.py \\\n"
        "    --review ./results/seeded_review_agentic.json \\\n"
        "    --ground-truth ./benchmarks/ground_truth.json \\\n"
        "    --report ./results/eval_agentic.json"
    )

    doc.add_heading("9.5 Fine-tune (QLoRA)", level=2)
    add_code(doc, "python training/fine_tune_go_reviewer.py --config training/training_config.yaml")

    doc.add_heading("9.6 Start the API server", level=2)
    add_code(doc, "uvicorn serving.api:app --host 0.0.0.0 --port 8000")

    # ---------- 10. Model architectures ----------
    doc.add_heading("10. Model Architectures", level=1)
    doc.add_paragraph(
        "The system uses three distinct model families: a generator LLM (the reviewer), "
        "an embedding model (for the RAG vector store), and the fine-tuning base model "
        "for the optional fine-tuned reviewer."
    )

    doc.add_heading("10.1 Generator: Qwen2.5-Coder-7B (default, via Ollama)", level=2)
    add_table(doc,
        ["Property", "Value"],
        [
            ["Vendor / family", "Alibaba — Qwen2.5 (Qwen2 architecture)"],
            ["Tag used", "qwen2.5-coder:7b"],
            ["Architecture", "Decoder-only Transformer (causal LM)"],
            ["Parameters", "7.6B"],
            ["Hidden size / layers / heads", "3584 / 28 / 28 (4 KV heads, GQA)"],
            ["Context length", "32,768 tokens"],
            ["Position encoding", "RoPE (rotary)"],
            ["Activation / norm", "SwiGLU FFN, RMSNorm"],
            ["Attention", "Grouped-Query Attention (GQA), Flash-Attn compatible"],
            ["Tokenizer", "BPE, ~152K vocab (Qwen2 tokenizer)"],
            ["Quantization (Ollama)", "GGUF Q4_K_M (~4.7 GB on disk)"],
            ["License", "Apache-2.0 / Tongyi Qianwen"],
            ["Strengths", "Code understanding, multi-language, instruction following"],
            ["Role", "Default reviewer in rag-only mode and inside agent nodes"],
        ])

    doc.add_heading("10.2 Generator: DeepSeek-Coder-7B-Instruct-v1.5 (fine-tune base)", level=2)
    add_table(doc,
        ["Property", "Value"],
        [
            ["Vendor / family", "DeepSeek-AI"],
            ["HF id", "deepseek-ai/deepseek-coder-7b-instruct-v1.5"],
            ["Architecture", "Decoder-only Transformer (LLaMA-family)"],
            ["Parameters", "6.7B"],
            ["Hidden size / layers / heads", "4096 / 32 / 32"],
            ["Context length", "16,384 tokens (extended via RoPE scaling)"],
            ["Position encoding", "RoPE"],
            ["Activation / norm", "SwiGLU FFN, RMSNorm"],
            ["Attention", "MHA / GQA (Flash-Attn 2 supported)"],
            ["Tokenizer", "BPE, ~32K vocab"],
            ["Pretraining mix", "Code-heavy (multi-language) + general text"],
            ["Role", "Base for QLoRA fine-tuning -> ./go-reviewer-final"],
        ])

    doc.add_heading("10.3 Fine-tuned reviewer (QLoRA on top of DeepSeek-Coder-7B)", level=2)
    add_table(doc,
        ["Property", "Value"],
        [
            ["Method", "QLoRA — 4-bit NF4 base + LoRA adapters on FP16"],
            ["LoRA rank / alpha / dropout", "r=64, alpha=128, dropout=0.05"],
            ["LoRA target modules", "q_proj, k_proj, v_proj, o_proj, gate_proj, up_proj, down_proj"],
            ["Quantization", "bitsandbytes NF4, double-quant, bf16 compute"],
            ["Training framework", "transformers + trl (SFTTrainer)"],
            ["Schedule", "3 epochs, batch=2, grad_accum=4, lr=2e-4 cosine, 5% warmup"],
            ["Sequence length", "2048 tokens"],
            ["Output artifacts", "LoRA adapter (output_dir/) + merged FP16 model (final_model_dir/)"],
            ["Role", "Reviewer in fine-tune-only and hybrid modes"],
        ])

    doc.add_heading("10.4 Embedding model: nomic-embed-text-v1.5", level=2)
    add_table(doc,
        ["Property", "Value"],
        [
            ["Vendor / family", "Nomic AI"],
            ["HF id", "nomic-ai/nomic-embed-text-v1.5"],
            ["Architecture", "BERT-style encoder (Nomic-BERT)"],
            ["Parameters", "~137M"],
            ["Hidden size / layers / heads", "768 / 12 / 12"],
            ["Max sequence length", "8192 tokens (Matryoshka, supports 2048/4096/8192)"],
            ["Output dimension", "768 (default; supports 64–768 via Matryoshka)"],
            ["Distance metric", "Cosine"],
            ["Pooling", "Mean pooling over tokens"],
            ["Role", "Embeds standards docs and code snippets for Qdrant retrieval"],
            ["License", "Apache-2.0"],
        ])

    doc.add_heading("10.5 Vector store: Qdrant (local, embedded mode)", level=2)
    add_table(doc,
        ["Property", "Value"],
        [
            ["Backend", "Qdrant local (file-backed, no separate server)"],
            ["Collection", "go_coding_standards"],
            ["Vector dim", "768"],
            ["Distance metric", "Cosine"],
            ["Index type", "HNSW (default Qdrant config)"],
            ["Storage path", "rag/qdrant_db/"],
            ["Source documents", "standards/*.md + standards/rules.json"],
            ["Chunker", "MarkdownTextSplitter (500 tokens, 50 overlap)"],
        ])

    doc.add_heading("10.6 How the models compose at runtime", level=2)
    add_code(doc, (
        "code chunk ──▶ nomic-embed-text-v1.5 ──▶ vector(768)\n"
        "                                            │\n"
        "                                            ▼\n"
        "                                       Qdrant (HNSW, cosine)\n"
        "                                            │\n"
        "                                            ▼\n"
        "                                  top-k rules (Documents)\n"
        "                                            │\n"
        "                                            ▼\n"
        "system + rules + code  ──▶  Qwen2.5-Coder-7B (Ollama)            ─┐\n"
        "                       OR  ─▶  DeepSeek-Coder-7B + LoRA (local)  ─┤──▶ findings (text)\n"
        "                       OR  ─▶  merged fine-tuned model (local)   ─┘\n"
        "                                            │\n"
        "                                            ▼\n"
        "                                    parse_findings()\n"
    ))

    # ---------- 11. Simple RAG vs Agentic RAG (architecture) ----------
    doc.add_heading("11. Simple RAG vs Agentic RAG: Architecture Comparison", level=1)
    doc.add_paragraph(
        "Both modes use the same retriever (Qdrant + nomic-embed-text-v1.5) and the "
        "same generator (Ollama or local fine-tuned model). They differ in how many "
        "LLM calls are made per chunk, how context is gathered, and how findings are "
        "produced and merged."
    )

    doc.add_heading("10.1 Simple RAG pipeline (single-shot retrieval)", level=2)
    add_code(doc, (
        "for each GoChunk:\n"
        "    rules     = retriever.retrieve(chunk.code, top_k=5)\n"
        "    prompt    = SYSTEM + format_rules(rules) + chunk.code\n"
        "    raw       = generate(prompt)        # 1 LLM call\n"
        "    findings += parse_findings(raw)\n"
        "\n"
        "deduplicate -> rank -> RepositoryReport"
    ))
    add_bullets(doc, [
        "Exactly 1 LLM call per chunk.",
        "Single retrieval call before generation; no second pass.",
        "Prompt = system instructions + top-k retrieved rules + full chunk code.",
        "No tool execution, no AST inspection, no lint runs.",
        "Cheapest and fastest; most predictable.",
    ])

    doc.add_heading("10.2 Agentic RAG pipeline (LangGraph multi-agent)", level=2)
    add_code(doc, (
        "for each GoChunk:\n"
        "    state     = ReviewState(code, file_path, repo_path, findings=[])\n"
        "    state     = planner_node(state)         # 1 LLM call: select 2-5 agents\n"
        "\n"
        "    for agent in selected_agents:           # parallel branches\n"
        "        rules        = retriever.retrieve(code, agent.rag_category, top_k=5)\n"
        "        tool_outputs = run_tools(agent.tools, repo_path)   # deterministic\n"
        "        prompt       = persona + focus + rules + tool_outputs + code\n"
        "        agent_raw    = generate(prompt)     # 1 LLM call per agent\n"
        "        state.findings += parse_findings(agent_raw)\n"
        "\n"
        "    state = aggregator_node(state)           # dedup + rank across agents\n"
        "\n"
        "RepositoryReport"
    ))
    add_bullets(doc, [
        "Planner makes 1 LLM call to pick 2–5 specialist agents.",
        "Each selected agent makes 1 LLM call → typically 3–6 LLM calls per chunk.",
        "Tools (AST parse, regex search, dependency graph, golangci-lint, go vet, schema query) execute in Python; outputs are materialized into the prompt up front.",
        "Each agent retrieves rules filtered by its rag_category (security / concurrency / performance / logging / architecture / database).",
        "Aggregator deduplicates findings across agents.",
    ])

    doc.add_heading("10.3 Side-by-side architecture", level=2)
    add_table(doc,
        ["Aspect", "Simple RAG", "Agentic RAG"],
        [
            ["LLM calls per chunk", "1", "3–6 (planner + selected agents)"],
            ["Retrieval calls per chunk", "1 (general)", "1 per selected agent (category-filtered)"],
            ["Tool calls per chunk", "0", "0–N per agent (AST, search, lint, deps, schema)"],
            ["Prompt context", "code + top-k rules", "code + filtered rules + tool outputs + persona/focus"],
            ["Concurrency", "Sequential per chunk", "Agents run in parallel (LangGraph fan-out)"],
            ["Findings merge", "Single LLM output → parse", "Multi-agent findings → aggregator → dedup"],
            ["Determinism", "Highest (one decoding pass)", "High (still deterministic per agent), but more variance surface"],
            ["Latency / chunk", "1× generation", "~3–6× generation + tool I/O"],
            ["Token cost", "Low", "3–6× higher"],
            ["Strengths", "Fast, predictable, cheap", "Domain expertise, rich evidence, parallel coverage"],
            ["Weaknesses", "One model must be a generalist", "Specialists can over-fire on their category"],
            ["Best for", "Small/medium repos, frequent runs, CI gating", "Large repos, deep audits, when tool evidence matters"],
        ])

    doc.add_heading("10.4 Why agentic underperformed on this benchmark", level=2)
    add_bullets(doc, [
        "Seeded benchmark files contain a single, isolated violation each — there is no extra repo context for tools to surface, so tool outputs add noise without new signal.",
        "Each specialist agent is biased toward its own category; the security agent over-reported SEC-001 on files where the true issue was EH-002 / SEC-003.",
        "With 3–6 LLM calls per chunk, every call is a chance to introduce a false positive that survives dedup.",
        "Simple RAG benefits when the model can see all relevant rules at once and decide globally — single decision, single output.",
        "Agentic mode pays off on real codebases where AST/dep-graph/lint evidence is non-trivial, multi-file context matters, and category specialists can raise issues that a generalist would miss.",
    ])

    # ---------- 12. Evaluation results ----------
    doc.add_heading("12. Evaluation Results: Simple RAG vs Agentic RAG", level=1)
    doc.add_paragraph(
        "Both pipelines were run on the seeded benchmark "
        "(benchmarks/seeded_repo, 17 violations + 1 clean file) using "
        "qwen2.5-coder:7b via Ollama and graded against ground_truth.json. "
        "Source data: results/seeded_eval_report.json (simple RAG) and "
        "results/eval_agentic.json (agentic RAG)."
    )

    doc.add_heading("12.1 Aggregate metrics", level=2)
    add_table(doc,
        ["Metric", "Simple RAG", "Agentic RAG", "Delta"],
        [
            ["Files", "18", "18", "—"],
            ["True positives", "9", "8", "-1"],
            ["False positives", "1", "2", "+1"],
            ["False negatives", "8", "9", "+1"],
            ["True negatives", "288", "287", "-1"],
            ["Precision", "0.90", "0.80", "-0.10"],
            ["Recall", "0.5294", "0.4706", "-0.0588"],
            ["F1 score", "0.6667", "0.5926", "-0.0741"],
            ["Accuracy", "0.9706", "0.9641", "-0.0065"],
            ["File exact-match", "0.50", "0.50", "0.00"],
            ["False positive rate", "0.10", "0.20", "+0.10"],
            ["Clean-file accuracy", "1.00", "1.00", "0.00"],
        ])
    doc.add_paragraph(
        "On this benchmark and model, the simple single-shot RAG actually outperformed "
        "the agentic multi-agent pipeline on every aggregate metric. The agentic mode "
        "produced two false positives (both predicting SEC-001 on files that were "
        "actually SEC-003 and EH-002 violations) and missed three rules that simple "
        "RAG caught (EH-001, EH-003, CTX-002)."
    )

    doc.add_heading("12.2 Per-rule comparison", level=2)
    add_table(doc,
        ["Rule", "Simple TP/FP/FN", "Agentic TP/FP/FN", "Notes"],
        [
            ["SEC-001", "1 / 0 / 0", "1 / 2 / 0", "Agentic over-fires SEC-001 on other files"],
            ["SEC-002", "0 / 0 / 1", "0 / 0 / 1", "Both miss SQL injection"],
            ["SEC-003", "0 / 0 / 1", "0 / 0 / 1", "Both miss JWT secret"],
            ["EH-001", "1 / 0 / 0", "0 / 0 / 1", "Simple wins"],
            ["EH-002", "1 / 0 / 0", "0 / 0 / 1", "Simple wins"],
            ["EH-003", "1 / 0 / 0", "0 / 0 / 1", "Simple wins"],
            ["CTX-001", "0 / 0 / 1", "0 / 0 / 1", "Both miss"],
            ["CTX-002", "1 / 0 / 0", "0 / 0 / 1", "Simple wins"],
            ["LOG-001", "1 / 0 / 0", "1 / 0 / 0", "Tie"],
            ["LOG-002", "0 / 0 / 1", "1 / 0 / 0", "Agentic wins"],
            ["NAM-001", "0 / 1 / 1", "1 / 0 / 0", "Agentic wins"],
            ["NAM-002", "0 / 0 / 1", "0 / 0 / 1", "Both miss"],
            ["CONC-001", "0 / 0 / 1", "0 / 0 / 1", "Both miss goroutine leak"],
            ["CONC-002", "1 / 0 / 0", "1 / 0 / 0", "Tie"],
            ["PERF-001", "1 / 0 / 0", "1 / 0 / 0", "Tie"],
            ["DOC-001", "1 / 0 / 0", "1 / 0 / 0", "Tie"],
            ["TEST-001", "1 / 0 / 0", "1 / 0 / 0", "Tie"],
        ])

    doc.add_heading("12.3 Per-file outcomes", level=2)
    add_table(doc,
        ["File", "Expected", "Simple RAG", "Agentic RAG"],
        [
            ["sec001_hardcoded.go", "SEC-001", "PASS", "PASS"],
            ["sec002_sqli.go", "SEC-002", "MISS", "MISS"],
            ["sec003_jwt_secret.go", "SEC-003", "MISS", "PARTIAL (predicted SEC-001)"],
            ["eh001_bare_return.go", "EH-001", "PASS", "MISS"],
            ["eh002_ignored_error.go", "EH-002", "PASS", "PARTIAL (predicted SEC-001)"],
            ["eh003_panic.go", "EH-003", "PASS", "MISS"],
            ["ctx001_ctx_not_first.go", "CTX-001", "MISS", "MISS"],
            ["ctx002_ctx_in_struct.go", "CTX-002", "PASS", "MISS"],
            ["log001_println.go", "LOG-001", "PASS", "PASS"],
            ["log002_printf.go", "LOG-002", "MISS", "PASS"],
            ["nam001_underscore.go", "NAM-001", "MISS", "PASS"],
            ["nam002_acronym.go", "NAM-002", "MISS", "MISS"],
            ["conc001_goroutine.go", "CONC-001", "MISS", "MISS"],
            ["conc002_bidirectional.go", "CONC-002", "PASS", "PASS"],
            ["test001_no_table.go", "TEST-001", "PASS", "PASS"],
            ["perf001_slice_alloc.go", "PERF-001", "NOISE (extra NAM-001)", "PASS"],
            ["doc001_no_comment.go", "DOC-001", "PASS", "PASS"],
            ["clean_code.go", "—", "PASS (no FP)", "PASS (no FP)"],
        ])

    doc.add_heading("12.4 Observations", level=2)
    add_bullets(doc, [
        "Both modes detect easy single-pattern rules (LOG-001, PERF-001, DOC-001, TEST-001, SEC-001, CONC-002) reliably.",
        "Both modes consistently miss SEC-002 (SQL injection), SEC-003 (JWT secret), CTX-001, CONC-001, NAM-002 — these need either better retrieval, more context, or a fine-tuned model.",
        "Simple RAG had higher recall on the error-handling family (EH-001, EH-002, EH-003) and CTX-002.",
        "Agentic RAG had higher recall on naming + observability (NAM-001, LOG-002) and avoided the spurious NAM-001 noise on perf001.",
        "Agentic mode's main failure mode here was the security agent over-reporting SEC-001 on unrelated files — a precision regression.",
        "Net: on this 18-file benchmark with this model, simple RAG is the better operating point. The agentic pipeline becomes more valuable as the codebase grows and tool outputs (AST, dependency graph, lint) materially expand the per-chunk context.",
    ])

    # ---------- 13. Approach comparison ----------
    doc.add_heading("13. Approach Comparison: RAG-Only vs Fine-Tuned vs Hybrid", level=1)
    doc.add_paragraph(
        "The system supports three deployment shapes. The choice should be driven by "
        "rule-change frequency, available labeled data, latency budget, and operational "
        "complexity tolerance."
    )
    add_table(doc,
        ["Evaluation Area", "RAG-Only (Qwen via Ollama)", "Fine-Tuned Only", "Hybrid (RAG + FT)"],
        [
            ["Rule freshness", "High — uses latest standards from vector DB at inference time.",
             "Low–Medium — requires retraining when standards change.",
             "High — uses fresh retrieved rules and learned behavior."],
            ["Consistency of findings", "Medium — depends on retrieval quality and prompt construction.",
             "High — model behavior is stable after training.",
             "High — most consistent when retrieval and model are both strong."],
            ["Setup complexity", "Low–Medium — build vector store and run Ollama.",
             "High — data curation, training infra, model artifacts.",
             "Highest — needs both training and retrieval pipelines."],
            ["Iteration speed", "Fast for rule updates; no retraining needed.",
             "Slow for rule updates; retraining cycle required.",
             "Medium — rule edits are fast, behavior changes still need retraining."],
            ["Runtime latency", "Medium — retrieval + generation per chunk.",
             "Fastest if model is local and optimized.",
             "Medium–High — retrieval overhead plus generation."],
            ["Data requirement", "Low — strong standards docs are enough to start.",
             "High — needs high-quality labeled examples.",
             "High — needs both strong docs and labeled examples."],
            ["Quality ceiling", "Medium–High with excellent standards and prompts.",
             "High for recurring organization-specific patterns.",
             "Highest in production when tuned correctly."],
            ["Best use case", "Early rollout, fast validation, frequent rule changes.",
             "Stable standards, high-volume reviews, strict consistency.",
             "Mature production system with continuous quality targets."],
        ])

    doc.add_heading("13.1 Practical notes for this setup", level=2)
    add_bullets(doc, [
        "Current deployment uses RAG-only with --ollama-model qwen2.5-coder:7b — ideal for rapid iteration.",
        "Fine-tuned-only and Hybrid require a trained model directory (e.g. ./go-reviewer-final).",
        "If standards change frequently, keep RAG active even after fine-tuning.",
        "Use Hybrid when the team needs both rule freshness and highly consistent classification output.",
    ])

    doc.add_heading("13.2 Recommended rollout path", level=2)
    add_bullets(doc, [
        "Start with RAG-only (simple) for immediate value and standards validation.",
        "Curate training data from accepted findings and human-reviewed corrections.",
        "Move to Hybrid in production once F1 on the seeded benchmark stabilizes above the target threshold.",
    ])

    # ---------- 14. Consistency & determinism ----------
    doc.add_heading("14. Consistency & Determinism", level=1)
    doc.add_paragraph(
        "LLMs are non-deterministic by default due to probabilistic token sampling. "
        "The same prompt can produce different outputs across runs. This section "
        "documents the techniques applied to make this pipeline reproducible and the "
        "empirical verification."
    )

    doc.add_heading("14.1 Observed problem", level=2)
    doc.add_paragraph(
        "Two consecutive runs with the default sampling configuration "
        "(temperature=0.3, top_p=0.95, no fixed seed) produced different findings on "
        "the same Go repository — different counts and different severity mixes."
    )

    doc.add_heading("14.2 Why prompt instructions alone do not work", level=2)
    doc.add_paragraph(
        "Adding 'be consistent' to the prompt does not fix the underlying randomness. "
        "Variability comes from the token sampling math, not from the model ignoring "
        "instructions. Prompt engineering helps quality but cannot enforce byte-identical "
        "outputs."
    )

    doc.add_heading("14.3 Techniques applied", level=2)
    add_table(doc,
        ["Technique", "Effect"],
        [
            ["temperature = 0.0", "Forces greedy decoding; the model always picks the most likely next token."],
            ["top_p = 1.0", "Disables nucleus sampling so filtering does not introduce randomness."],
            ["seed = 42", "Fixes Ollama's random state so any remaining stochastic step is reproducible."],
            ["Pinned model tag", "qwen2.5-coder:7b explicitly; avoid silent upgrades."],
            ["Frozen vector store", "Don't rebuild between comparison runs."],
        ])

    doc.add_heading("14.4 Empirical verification", level=2)
    doc.add_paragraph(
        "After applying temperature=0, top_p=1.0, seed=42, the pipeline was run twice "
        "with identical inputs and outputs were compared programmatically."
    )
    add_table(doc,
        ["Metric", "Run A (test_deterministic.json)", "Run B (test_verify.json)"],
        [
            ["Total findings", "8", "8"],
            ["Critical", "2", "2"],
            ["High", "5", "5"],
            ["Medium", "1", "1"],
            ["Rule IDs + file + line match", "Yes", "Yes"],
        ])
    doc.add_paragraph(
        "Result: identical findings in count, severity, rule IDs, files, and line numbers. "
        "Reproducibility is confirmed for this setup."
    )

    doc.add_heading("14.5 Consistency by approach", level=2)
    add_table(doc,
        ["Approach", "Default consistency", "With deterministic settings", "Notes"],
        [
            ["RAG-Only (Ollama)", "Low–Medium", "High",
             "Qdrant local mode is deterministic for the same query and frozen index."],
            ["Fine-Tuned Only", "Medium–High", "Highest",
             "Less prompt-sensitive, but still needs deterministic decoding."],
            ["Hybrid", "Medium", "High",
             "Rebuilding the vector store changes retrieved context; rebuild intentionally."],
        ])

    doc.add_heading("14.6 Recommended deterministic configuration", level=2)
    add_bullets(doc, [
        "Set temperature to 0.0 (greedy decoding).",
        "Set top_p to 1.0 so nucleus sampling is a no-op.",
        "Set a fixed seed (e.g. 42) and keep it in version control.",
        "Pin the model tag (qwen2.5-coder:7b) and avoid upgrading silently.",
        "Freeze the vector store; do not rebuild between comparison runs.",
        "Record model version, seed, and config hash in every report for auditability.",
    ])
    doc.add_paragraph(
        "Caveat: exact byte-level reproducibility also depends on Ollama version, "
        "hardware (CPU vs GPU), and numerical precision. Day-to-day runs on the same "
        "machine with the same model tag should be effectively identical after applying "
        "these settings."
    )

    # ---------- 15. Operational notes ----------
    doc.add_heading("15. Operational Notes", level=1)
    add_bullets(doc, [
        "Proxy / localhost: pipeline talks to Ollama at http://localhost:11434. If HTTP_PROXY/HTTPS_PROXY are set, ensure no_proxy includes both 'localhost' and '127.0.0.1'.",
        "Temp dir permissions: go_parser.py compiles ast_helper.go in tempfile.gettempdir(). On shared hosts, set TMPDIR to a user-writable directory (e.g. $HOME/.cache/go_reviewer_tmp).",
        "HF cache: embedding model downloads under HF_HOME / TRANSFORMERS_CACHE. Permission-denied cache warnings are non-fatal but indicate a shared cache; point HF_HOME at a user-owned dir to silence them.",
        "Determinism: set temperature=0 and a fixed seed in ReviewConfig for reproducible output. Pin the Ollama model tag and avoid silent upgrades.",
        "Vector store stability: don't rebuild the vector store between comparison runs unless intentional — retrieved context affects findings.",
    ])

    # ---------- 16. Component cheat sheet ----------
    doc.add_heading("16. Component Cheat Sheet", level=1)
    add_table(doc,
        ["Component", "Purpose", "Input", "Output"],
        [
            ["pipeline/review_pipeline.py", "Orchestrator", "repo / snippet", "RepositoryReport"],
            ["pipeline/go_parser.py", "Go AST chunking", "repo path", "List[GoChunk]"],
            ["pipeline/deduplication.py", "Parse + dedup findings", "LLM text", "List[dict]"],
            ["pipeline/report_generator.py", "Format reports", "RepositoryReport", "Markdown / JSON / SARIF"],
            ["rag/retriever.py", "Vector retrieval", "code snippet", "List[Document]"],
            ["rag/build_vector_store.py", "Index standards", "markdown + JSON", "Qdrant DB"],
            ["rag/agent_tools.py", "Deterministic tools", "repo path", "dict of callables"],
            ["rag/langgraph_agents.py", "Multi-agent reviewer", "code + repo path", "List[dict]"],
            ["training/fine_tune_go_reviewer.py", "QLoRA fine-tune", "JSONL data", "LoRA adapter + merged model"],
            ["training/evaluate_seeded.py", "Seeded benchmark eval", "review + ground truth", "metrics report"],
            ["training/evaluate.py", "Held-out eval", "model + test JSONL", "per-category metrics"],
            ["dataset/build_dataset.py", "Build training data", "linter output / repos", "JSONL"],
            ["cli/review.py", "CLI front-end", "CLI args", "review output"],
            ["serving/api.py", "REST API", "HTTP requests", "JSON responses"],
        ])

    doc.save(OUTPUT_PATH)
    print(f"[OK] wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
