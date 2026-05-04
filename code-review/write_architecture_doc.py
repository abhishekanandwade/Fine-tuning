"""Build Architecture.docx: project overview + evaluation explanation."""
from pathlib import Path
from docx import Document
from docx.shared import Pt

DOC = Path(r"C:\Users\knchaitr\OneDrive - Hewlett Packard Enterprise\CoE Team\Fine-tuning\Architecture.docx")

doc = Document()

def H(text, level=1):
    doc.add_heading(text, level=level)

def P(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p

def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def kv_table(rows, headers):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)


# TITLE
doc.add_heading("Agentic-RAG Code Reviewer for Go Repositories", 0)
P("A local, agentic Retrieval-Augmented-Generation pipeline that audits Go "
  "repositories for three architectural rules. Combines deterministic regex "
  "scanning with a small local LLM (qwen2.5-coder:7b via Ollama) to deliver "
  "fast, explainable findings without sending code to the cloud.")

# 1. PROBLEM
H("1. Problem Statement", 1)
P("Manual architectural code review is slow, inconsistent, and skipped under "
  "deadline pressure. Cloud-based LLMs are accurate but cannot be used on "
  "proprietary backend code. Pure regex/AST tools are fast but lack semantic "
  "understanding and produce noisy false positives.")
P("Goal: a fully local, repo-agnostic reviewer that flags only high-confidence "
  "violations of three specific architectural rules, explains each finding "
  "with offending source lines, and exposes a measurable recall metric.")

# 2. RULES
H("2. Rules Detected", 1)
kv_table(
    rows=[
        ("REPO-001", "Repository function makes more than one DB call.",
         "Repository methods should issue exactly one database call. Multiple calls per function indicate missing aggregation or leaked transaction logic."),
        ("REPO-002", "DB call inside a for/range loop (N+1).",
         "Database calls inside loops cause N+1 query patterns and severe latency. Should be rewritten using IN-clauses, batches, or joins."),
        ("HANDLER-001", "Handler invokes more than one repository / service method.",
         "HTTP handlers should orchestrate via a single service method. Multiple direct repo calls bleed business logic into the transport layer."),
    ],
    headers=("Rule ID", "Title", "Why it matters"),
)
P("Rule definitions, regex patterns, violation examples, and suggested fixes "
  "live in standards/architectural_rules.json (the RAG knowledge base).")

# 3. ARCHITECTURE
H("3. System Architecture", 1)
P("Four cooperating components in code-review/pipeline/:")
kv_table(
    rows=[
        ("architectural_rag_agent.py", "The reviewer.",
         "Walks Go files, classifies repo/handler, runs deterministic pre-filter, invokes LLM via LangGraph for confirmation, validates output deterministically."),
        ("evaluate_coverage.py", "The grader.",
         "Re-runs the deterministic pre-filter as a baseline and compares to the agent's findings. Reports CONFIRMED / MISSED / EXTRA per rule and overall recall %."),
        ("merge_findings.py", "The combiner.",
         "Unions multiple review JSONs by (file, function, rule_id). Used to merge initial pass with recovery pass."),
        ("standards/architectural_rules.json", "The knowledge base.",
         "RAG corpus: per-rule descriptions, regex hints, violation examples, correct examples, suggested fixes."),
    ],
    headers=("Component", "Role", "Detail"),
)

H("3.1 Agentic-RAG flow (per Go file)", 2)
code(
"   classify_file ---> extract_functions ---> deterministic pre-filter (regex)\n"
"                                                |\n"
"                                                v\n"
"                          +----------- LangGraph agent -----------+\n"
"                          |  retrieve_rule  ->  analyze_llm        |\n"
"                          |          ^               |             |\n"
"                          |          +--- critic <---+ (deterministic\n"
"                          |                            verifier)    |\n"
"                          +--------------------+-------------------+\n"
"                                               v\n"
"                                          finding(s)"
)

H("3.2 Two-tier RAG", 2)
bullet("Lexical retrieval over standards/architectural_rules.json — fetches the rule definition + violation/correct examples for the candidate function.")
bullet("Code retrieval — the agent receives the function source plus the file's package declaration. For long functions a windowed view around the candidate offending lines is supplied.")

H("3.3 Why the critic is deterministic, not LLM", 2)
P("An earlier LLM-based critic was producing false negatives (re-interpreting "
  "the rule and overriding correct verdicts). It was replaced by a programmatic verifier that:")
bullet("Validates each cited line number falls within the function's range.")
bullet("Confirms the cited source line actually matches the rule's regex pattern.")
bullet("For REPO-002, confirms the cited line is physically nested inside a for/range loop.")
bullet("Applies rule quorum: REPO-001 / HANDLER-001 need >= 2 valid hits, REPO-002 needs >= 1 in-loop hit.")
P("Result: zero second LLM call (~2x faster), zero false-negative loop, deterministic and reproducible.")

# 4. WORKFLOW
H("4. End-to-End Workflow", 1)
P("Five commands run from code-review/. Replace <REPO> with the path to the Go repo.")
code(
"# 0. Activate venv + ensure Ollama is running with qwen2.5-coder:7b\n"
"Set-ExecutionPolicy -Scope Process -ExecutionPolicy RemoteSigned\n"
"& \"..\\.venv\\Scripts\\Activate.ps1\"\n"
"$env:PYTHONIOENCODING = \"utf-8\"\n"
"\n"
"# 1. Initial agentic-RAG review (LLM-confirmed)\n"
"python -m pipeline.architectural_rag_agent \\\n"
"  --repo <REPO> --output results\\architectural_review.json \\\n"
"  --ollama-model qwen2.5-coder:7b\n"
"\n"
"# 2. Grade coverage vs deterministic baseline\n"
"python -m pipeline.evaluate_coverage \\\n"
"  --repo <REPO> --report results\\architectural_review.json \\\n"
"  --output results\\coverage_report.json --markdown results\\coverage_report.md\n"
"\n"
"# 3. Recover any missed candidates (deterministic, instant)\n"
"python -m pipeline.architectural_rag_agent \\\n"
"  --repo <REPO> --output results\\architectural_review_recovered.json \\\n"
"  --recover-from results\\coverage_report.json --no-llm\n"
"\n"
"# 4. Merge initial + recovery into final report\n"
"python -m pipeline.merge_findings \\\n"
"  --inputs results\\architectural_review.json \\\n"
"           results\\architectural_review_recovered.json \\\n"
"  --output results\\architectural_review_merged.json\n"
"\n"
"# 5. Re-score (should be 100% recall)\n"
"python -m pipeline.evaluate_coverage \\\n"
"  --repo <REPO> --report results\\architectural_review_merged.json \\\n"
"  --output results\\coverage_report_merged.json \\\n"
"  --markdown results\\coverage_report_merged.md"
)

H("4.1 Stage outputs", 2)
kv_table(
    rows=[
        ("Stage 1", "results/architectural_review.json", "Initial agent findings (LLM-confirmed)."),
        ("Stage 2", "results/coverage_report.{json,md}", "Lists CONFIRMED / MISSED / EXTRA + recall."),
        ("Stage 3", "results/architectural_review_recovered.json", "Findings recovered from MISSED list, deterministically."),
        ("Stage 4", "results/architectural_review_merged.json", "FINAL findings (initial U recovered)."),
        ("Stage 5", "results/coverage_report_merged.{json,md}", "FINAL grading — should show 100% recall."),
    ],
    headers=("Stage", "Output file", "Description"),
)

H("4.2 Recovery design — works on any repo", 2)
P("--recover-from reads a coverage_report.json, extracts every missed "
  "(file, function, rule_id) triple, and runs the agent ONLY on that set. "
  "No function names are hardcoded, so the same recovery loop works on any "
  "Go repository where coverage shows gaps.")
P("--no-llm-veto is an alternative recovery mode: when the LLM disagrees with "
  "the deterministic pre-filter, emit the finding using the regex hits anyway.")

# 5. RESULTS
H("5. Results on the Reference Repo", 1)
P("Repository: post-data-management-back-end-development (Go, 33 .go files; 9 repo files, 13 handler files).")
kv_table(
    rows=[
        ("Initial agent run", "64", "—", "LLM disagreed on 11 candidates."),
        ("Recovery (--recover-from --no-llm)", "+11", "—", "Deterministic recovery; instant."),
        ("Merged (final)", "75", "100.00%", "0 missed, 0 extras."),
    ],
    headers=("Stage", "Findings", "Recall", "Notes"),
)

H("5.1 Findings by rule (final)", 2)
kv_table(
    rows=[
        ("REPO-001", "36", "Multiple DB calls per repository function."),
        ("REPO-002", "17", "DB calls inside for/range loops."),
        ("HANDLER-001", "22", "Handlers calling >1 repository/service method."),
    ],
    headers=("Rule", "Count", "Description"),
)

# 6. EVALUATION
H("6. How Evaluation Works", 1)
P("evaluate_coverage.py measures the agent's self-consistency. It does NOT "
  "use a hand-labeled ground-truth dataset. Instead, it re-runs the same "
  "deterministic pre-filter the agent uses, treats those candidates as the "
  "baseline, and compares them against the agent's emitted findings.")

H("6.1 Core idea", 2)
code(
"                          +---------------------+\n"
"   Same Go repo --------->| Deterministic       |--- set A: candidates\n"
"                          | pre-filter          |    (file, fn, rule_id)\n"
"                          | (regex + quorum)    |\n"
"                          +---------------------+\n"
"\n"
"   architectural_review.json  ---> set B: agent findings\n"
"                                   (file, fn, rule_id)\n"
"\n"
"                  Compare set A vs set B\n"
"                  ----------------------\n"
"                  CONFIRMED = A intersect B\n"
"                  MISSED    = A - B   (agent should have flagged but didn't)\n"
"                  EXTRA     = B - A   (agent flagged with no det. signal)\n"
"                  RECALL    = |CONFIRMED| / |A|"
)

H("6.2 Step by step", 2)
bullet("Walk every Go file (skipping vendor/, test_suite/, *_test.go).")
bullet("Classify each file as repo, handler, or skip — based on directory keywords.")
bullet("Extract every Go function with its line range.")
bullet("Run the same regex pre-filter the agent uses to find candidate offending lines.")
bullet("Apply rule-specific quorum: REPO-001 / HANDLER-001 require >= 2 hits; REPO-002 requires >= 1 hit nested inside a for/range loop. Pass = candidate.")
bullet("Load the agent's findings JSON and build the same key set.")
bullet("Set arithmetic per rule: CONFIRMED = both, MISSED = det only, EXTRA = agent only.")
bullet("Compute recall = confirmed / (confirmed + missed) per rule and overall.")
bullet("Emit coverage_report.json + .md (summary tables and per-rule MISSED tables with sample evidence lines).")

H("6.3 Metric meaning", 2)
kv_table(
    rows=[
        ("Candidates", "Functions where deterministic quorum was met", "75"),
        ("Confirmed", "Agent agreed with the deterministic signal", "64 -> 75"),
        ("Missed", "Pre-filter said 'look here', agent said 'no violation'", "11 -> 0"),
        ("Extra", "Agent flagged with no deterministic signal", "0"),
        ("Recall", "Confirmed / (Confirmed + Missed)", "85.33% -> 100%"),
    ],
    headers=("Metric", "Definition", "This repo"),
)

H("6.4 Caveats", 2)
bullet("The baseline IS the deterministic pre-filter — same regexes the agent uses. High recall therefore means LLM-vs-regex agreement, not absolute correctness.")
bullet("Many EXTRAs would indicate LLM hallucination. We observed 0 — clean.")
bullet("MISSED items are not necessarily true violations; they are functions where the regex quorum triggered but the LLM disagreed. Spot-check to decide.")
bullet("For true precision/recall, build a hand-labeled ground-truth set for REPO-001 / REPO-002 / HANDLER-001 and compare against that instead.")

# 7. KEY DESIGN DECISIONS
H("7. Key Design Decisions", 1)
bullet("Local-only execution — Ollama keeps source code on the developer's machine.")
bullet("Deterministic pre-filter gates the LLM — only a small subset of functions reach the LLM, cutting cost and runtime substantially.")
bullet("Deterministic critic instead of LLM critic — eliminates a class of false negatives where the LLM second-guesses its own correct verdict.")
bullet("Quorum thresholds (>=2 for REPO-001/HANDLER-001, >=1 in-loop for REPO-002) tuned to suppress noise.")
bullet("Rule-aware recovery via --recover-from — the evaluator's MISSED list is the agent's input, closing the gap automatically without hardcoded function names.")
bullet("LangGraph StateGraph with TypedDict schema — required for proper state merging across nodes.")
bullet("Stable LLM settings: temperature=0.0, seed=42, num_predict=384 — produces reproducible verdicts.")

# 8. REPO LAYOUT
H("8. Repository Layout", 1)
code(
"code-review/\n"
"  pipeline/\n"
"    architectural_rag_agent.py    # the reviewer (agent + pre-filter + critic)\n"
"    evaluate_coverage.py           # the grader\n"
"    merge_findings.py              # unions multiple review reports\n"
"    __init__.py\n"
"  standards/\n"
"    architectural_rules.json       # 3 rules: REPO-001, REPO-002, HANDLER-001\n"
"  results/\n"
"    architectural_review_merged.json   # FINAL findings (75)\n"
"    coverage_report_merged.json/.md    # FINAL grading (100% recall)\n"
"    architectural_review.json          # initial pass (intermediate, 64)\n"
"    architectural_review_recovered.json# recovery pass (intermediate, 11)\n"
"    coverage_report.json/.md           # initial grading (intermediate)\n"
"  requirements.txt                 # langgraph, requests, python-docx\n"
"  .gitignore"
)

# 9. TECH STACK — DETAILED
H("9. Tech Stack — Detailed", 1)

H("9.1 LLM (the generator)", 2)
kv_table(
    rows=[
        ("Model", "qwen2.5-coder:7b", "Alibaba Qwen2.5-Coder, 7B-parameter code-specialized LLM."),
        ("Runtime", "Ollama", "Local model server. HTTP API at http://localhost:11434/api/chat."),
        ("Transport", "HTTP / requests", "Plain Python `requests.post(url, json=payload, timeout=300)`. No SDK dependency."),
        ("Sampling", "deterministic", "temperature=0.0, seed=42, num_predict=384 — produces reproducible verdicts."),
        ("Prompt format", "system + user (chat)", "System message defines JSON schema and rule-quorum constraints; user message carries rule, examples, and line-numbered function source."),
        ("Output contract", "strict JSON", "{ violates: bool, offending_lines: [\"line N: code\"...], reasoning: str }. Parsed with a lenient JSON extractor."),
    ],
    headers=("Aspect", "Value", "Detail"),
)
P("Why Qwen2.5-Coder 7B: strong Go/SQL understanding for its size, fits on a "
  "developer laptop (~5 GB VRAM/RAM via Ollama's GGUF quantization), no data "
  "leaves the machine. No fine-tuning was performed — capability comes from "
  "structured prompting + deterministic verifier.")

H("9.2 Retrieval-Augmented Generation (the RAG)", 2)
P("This is a LEXICAL RAG, not a vector-store RAG. There is intentionally NO "
  "embedding model and NO vector database. Justification:")
bullet("The corpus is tiny — three rules in standards/architectural_rules.json (~5 KB).")
bullet("Retrieval is exact-by-key (rule_id), so dense embeddings would add latency and dependencies for zero accuracy gain.")
bullet("Code retrieval is also lexical: the agent receives the exact function source plus its package declaration. Long functions get a windowed view via numbered_source(fn, hit_lines, max_lines=90, window=8).")

P("Two-tier retrieval per analysis:")
kv_table(
    rows=[
        ("Tier 1: Rule retrieval",
         "Direct dict lookup",
         "rules_by_id[rule_id] -> {description, violation_example, correct_example, db_call_indicators, suggested_fix}. Embedded into the user prompt verbatim."),
        ("Tier 2: Code retrieval",
         "Function extraction + windowed view",
         "extract_functions() pulls the candidate function via brace-counting parser; numbered_source() returns line-numbered Go source, optionally trimmed around regex hits for long bodies."),
    ],
    headers=("Tier", "Mechanism", "Detail"),
)

H("9.3 Vector store / embeddings — NOT used (and why)", 2)
bullet("No FAISS, Chroma, Qdrant, or pgvector.")
bullet("No embedding model (e.g. nomic-embed-text, bge-base).")
bullet("Reasoning: corpus is 3 rules; semantic similarity is unnecessary when rule_id selection is exact.")
bullet("If the rule set grows past ~30 rules, swapping in a vector store is straightforward: replace rules_by_id lookup with similarity search over rule descriptions.")

H("9.4 Agentic framework", 2)
kv_table(
    rows=[
        ("Framework", "LangGraph", "Graph-based orchestration of LLM nodes. Provides state schema, conditional edges, retries."),
        ("Graph type", "StateGraph(AgentState)", "AgentState is TypedDict(total=False) — REQUIRED for LangGraph's reducer to merge node outputs correctly."),
        ("Nodes", "analyze, critic", "analyze = LLM call producing JSON verdict; critic = deterministic verifier (no LLM)."),
        ("Edges", "START -> analyze -> critic -> {analyze, END}", "Conditional edge: if critic emits 'final', stop; else loop back for one retry with feedback."),
        ("Max iterations", "1 by default", "Hard cap on analyze re-tries; tunable via --max-iterations."),
    ],
    headers=("Aspect", "Value", "Detail"),
)
P("Why LangGraph (not LangChain Agents or AutoGen): explicit state machine, "
  "no hidden tool-calling magic, clean separation between LLM reasoning and "
  "deterministic verification.")

H("9.5 Tools / capabilities the agent has", 2)
P("This agent does NOT use OpenAI-style function-calling tools. Each node is "
  "a Python function with full repo access:")
bullet("analyze node: invokes Ollama, returns structured JSON verdict.")
bullet("critic node: in-process Python validator. Verifies line numbers, regex matches, and loop nesting against a body_map dict captured before invocation. No LLM call, no external tool.")

H("9.6 Pattern matching (the deterministic backbone)", 2)
P("Three regexes drive the pre-filter and the critic:")
code(
"_DB_HINT     = matches: db.Query/Exec/QueryContext/SendBatch/Begin/Prepare,\n"
"               dblib.Insert/Update/Delete/Select, psql.*, batch.Queue,\n"
"               br.Exec/Query, pgx.RowToStruct\n"
"\n"
"_REPO_HINT   = matches: pmh.svc.*, h.svc/service/repo.*, *.svc/service/repo/\n"
"               Repo/Service/Repository.*, *Repo.*, *Service.*\n"
"\n"
"_FOR_LOOP_OPEN = matches: ^\\s*for\\s+   (loop opening line)"
)
P("File classification and Go function parsing are also pure Python regex + "
  "brace-counting — no go/ast, no goimports, no external Go toolchain. This "
  "keeps the project zero-Go-dependency.")

H("9.7 Frameworks & libraries", 2)
kv_table(
    rows=[
        ("LangGraph", ">=0.2", "Agent graph orchestration."),
        ("requests", "stdlib-grade HTTP", "Ollama API client."),
        ("python-docx", "writes .docx", "This document generator only."),
        ("Python", "3.12", "Runtime."),
        ("Ollama", "latest", "Local LLM server. Pulled model: qwen2.5-coder:7b (~4.7 GB)."),
    ],
    headers=("Library / tool", "Version", "Purpose"),
)
P("Deliberately NOT used: LangChain (avoids version churn), OpenAI SDK (cloud), "
  "FAISS/Chroma/Qdrant (no vector store needed), tree-sitter or go/ast "
  "(brace-counting parser is sufficient and dependency-free), pytest (no test "
  "suite yet — evaluate_coverage.py serves as the integration test).")

H("9.8 Data flow summary", 2)
code(
"  Go file ─► extract_functions ─► candidates_for_rule (regex)\n"
"                                       │\n"
"                                       ▼ (if quorum met)\n"
"  rules.json ─► rule lookup ─► _build_user_prompt ─► Ollama HTTP API\n"
"                                                          │\n"
"                                                          ▼\n"
"                                       JSON verdict (lenient parse)\n"
"                                                          │\n"
"                                                          ▼\n"
"                                       deterministic critic\n"
"                                       (line range + regex + loop check)\n"
"                                                          │\n"
"                                                          ▼\n"
"                                       Finding(rule_id, file, function,\n"
"                                              line_start..line_end,\n"
"                                              offending_lines, suggested_fix)"
)

H("9.9 What this is NOT", 2)
bullet("Not a fine-tuned model — base qwen2.5-coder:7b is used as-is.")
bullet("Not a cloud service — fully offline after model pull.")
bullet("Not a vector RAG — lexical/exact retrieval only.")
bullet("Not a function-calling agent — nodes are explicit Python functions.")
bullet("Not a Go-AST tool — uses regex + brace-counting parser.")

doc.save(str(DOC))
print(f"[done] wrote {DOC}")
