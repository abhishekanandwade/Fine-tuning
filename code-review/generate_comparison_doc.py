from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

output_path = "comparison_rag_vs_finetuned_qwen_ollama_v2.docx"

doc = Document()

# Title
p = doc.add_paragraph()
r = p.add_run("Comparison Report: RAG-Only vs Fine-Tuned vs Hybrid (RAG + Fine-Tuned)\nUsing Qwen via Ollama")
r.bold = True
r.font.size = Pt(16)
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

meta = doc.add_paragraph(
    f"Prepared on: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
    "Context: Go code review workflow using local Ollama-hosted Qwen models."
)
meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Objective
doc.add_heading("1. Objective", level=1)
doc.add_paragraph(
    "This document compares three deployment approaches for automated Go code review: "
    "RAG-only, Fine-tuned-only, and Hybrid (RAG + Fine-tuned). "
    "The comparison focuses on quality, maintainability, cost, performance, and operational complexity."
)

# Comparison table
doc.add_heading("2. Side-by-Side Comparison", level=1)
headers = [
    "Evaluation Area",
    "RAG-Only (Qwen via Ollama)",
    "Fine-Tuned Only (Qwen-derived)",
    "Hybrid (RAG + Fine-Tuned)",
]
rows = [
    (
        "Rule freshness",
        "High. Uses latest standards from vector DB at inference time.",
        "Low-Medium. Requires retraining when standards change.",
        "High. Uses fresh retrieved rules and learned behavior.",
    ),
    (
        "Consistency of findings",
        "Medium. Depends on retrieval quality and prompt construction.",
        "High. Model behavior is stable after training.",
        "High. Most consistent when retrieval and model are both strong.",
    ),
    (
        "Setup complexity",
        "Low-Medium. Build vector store and run Ollama.",
        "High. Requires data curation, training infra, model artifacts.",
        "Highest. Needs both training and retrieval pipelines.",
    ),
    (
        "Iteration speed",
        "Fast for rule updates. No retraining needed.",
        "Slow for rule updates. Retraining cycle required.",
        "Medium. Rule edits are fast, behavior changes still need retraining.",
    ),
    (
        "Runtime latency",
        "Medium. Retrieval + generation per chunk.",
        "Fastest if model is local and optimized.",
        "Medium-High. Retrieval overhead plus generation.",
    ),
    (
        "Data requirement",
        "Low. Strong standards docs are enough to start.",
        "High. Needs high-quality labeled examples.",
        "High. Needs both strong docs and labeled examples.",
    ),
    (
        "Quality ceiling",
        "Medium-High with excellent standards and prompts.",
        "High for recurring organization-specific patterns.",
        "Highest in production when tuned correctly.",
    ),
    (
        "Best use case",
        "Early rollout, fast validation, frequent rule changes.",
        "Stable standards, high-volume reviews, strict consistency.",
        "Mature production system with continuous quality targets.",
    ),
]

table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    hdr_cells[i].text = h

for row in rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text

# Practical notes
doc.add_heading("3. Practical Notes for Your Current Setup", level=1)
notes = [
    "You are currently running RAG-only with `--ollama-model qwen2.5-coder:7b`, which is ideal for rapid iteration.",
    "Fine-tuned-only and Hybrid require a trained model directory (for example `./go-reviewer-final`).",
    "If standards change frequently, keep RAG active even after fine-tuning.",
    "Use Hybrid when the team needs both rule freshness and highly consistent classification output.",
]
for n in notes:
    doc.add_paragraph(n, style="List Bullet")

# Recommendation
doc.add_heading("4. Recommendation", level=1)
doc.add_paragraph(
    "Recommended rollout path:\n"
    "1) Start with RAG-only for immediate value and standards validation.\n"
    "2) Build and curate training data from accepted findings.\n"
    "3) Move to Hybrid for production once model quality and evaluation metrics are stable."
)

# Consistency / Determinism Analysis
doc.add_heading("5. Consistency & Determinism Analysis", level=1)

doc.add_paragraph(
    "LLMs are non-deterministic by default due to probabilistic token sampling. "
    "The same prompt can produce different outputs across runs. This section documents "
    "the consistency problem, the techniques applied, and empirical results from the "
    "Qwen + Ollama setup."
)

doc.add_heading("5.1 Observed Problem", level=2)
doc.add_paragraph(
    "Two consecutive runs with the default sampling configuration (temperature=0.3, top_p=0.95, "
    "no fixed seed) produced different findings for the same Go repository:"
)
prob_rows = [
    ("Run 1 (test.json)", "9 findings", "Critical: 1, High: 5, Medium: 2, Info: 1"),
    ("Run 2 (test_1.json)", "Different count", "Different severity mix"),
]
t1 = doc.add_table(rows=1, cols=3)
t1.style = "Table Grid"
h1 = t1.rows[0].cells
h1[0].text = "Run"
h1[1].text = "Total Findings"
h1[2].text = "Severity Breakdown"
for r_ in prob_rows:
    c = t1.add_row().cells
    for i, t_ in enumerate(r_):
        c[i].text = t_

doc.add_heading("5.2 Why Prompt Instructions Alone Do Not Work", level=2)
doc.add_paragraph(
    "Adding instructions such as 'be consistent' or 'always produce the same answer' "
    "does not fix the underlying randomness. The variability comes from the token sampling "
    "math, not from the model ignoring instructions. Prompt engineering helps quality but "
    "cannot enforce byte-identical outputs."
)

doc.add_heading("5.3 Techniques Applied", level=2)
tech_rows = [
    ("temperature = 0.0", "Forces greedy decoding; the model always picks the most likely next token."),
    ("top_p = 1.0", "Disables nucleus sampling so filtering does not introduce randomness."),
    ("seed = 42", "Fixes Ollama's random state so any remaining stochastic step is reproducible."),
    ("Prompt 'be consistent'", "Not relied on. Useful for style but not for determinism."),
]
t2 = doc.add_table(rows=1, cols=2)
t2.style = "Table Grid"
h2 = t2.rows[0].cells
h2[0].text = "Technique"
h2[1].text = "Effect"
for r_ in tech_rows:
    c = t2.add_row().cells
    for i, t_ in enumerate(r_):
        c[i].text = t_

doc.add_heading("5.4 Empirical Verification", level=2)
doc.add_paragraph(
    "After applying temperature=0.0, top_p=1.0, and seed=42, the pipeline was run twice "
    "with identical inputs and the outputs were compared programmatically."
)
verify_rows = [
    ("Metric", "test_deterministic.json", "test_verify.json"),
    ("Total findings", "8", "8"),
    ("Critical", "2", "2"),
    ("High", "5", "5"),
    ("Medium", "1", "1"),
    ("Rule IDs + file + line match", "Yes", "Yes"),
]
t3 = doc.add_table(rows=0, cols=3)
t3.style = "Table Grid"
for r_ in verify_rows:
    c = t3.add_row().cells
    for i, t_ in enumerate(r_):
        c[i].text = t_

doc.add_paragraph(
    "Result: Both deterministic runs produced identical findings in count, severity, "
    "rule IDs, files, and line numbers. Reproducibility is confirmed for this setup."
)

doc.add_heading("5.5 Consistency Comparison Across Approaches", level=2)
cons_headers = ["Approach", "Default Consistency", "With Deterministic Settings", "Notes"]
cons_rows = [
    (
        "RAG-Only (Qwen via Ollama)",
        "Low-Medium. Output varies run to run.",
        "High. Identical outputs with temperature=0 and fixed seed.",
        "Retrieval is deterministic in Qdrant local mode for the same query and index.",
    ),
    (
        "Fine-Tuned Only",
        "Medium-High. Model bias toward learned patterns.",
        "Highest. Same decoding config makes output stable.",
        "Less prompt-sensitive, but still needs deterministic decoding.",
    ),
    (
        "Hybrid (RAG + Fine-Tuned)",
        "Medium. Adds retrieval variability on top of generation.",
        "High. Requires deterministic decoding and stable retrieval index.",
        "Rebuilding the vector store can change retrieved context; rebuild intentionally.",
    ),
]
t4 = doc.add_table(rows=1, cols=4)
t4.style = "Table Grid"
h4 = t4.rows[0].cells
for i, h in enumerate(cons_headers):
    h4[i].text = h
for r_ in cons_rows:
    c = t4.add_row().cells
    for i, t_ in enumerate(r_):
        c[i].text = t_

doc.add_heading("5.6 Recommended Deterministic Configuration", level=2)
doc.add_paragraph(
    "For reproducible code review runs across all three approaches with Ollama:"
)
rec_bullets = [
    "Set temperature to 0.0 (greedy decoding).",
    "Set top_p to 1.0 so nucleus sampling is a no-op.",
    "Set a fixed seed (for example 42) and keep it in version control.",
    "Pin the model tag (for example qwen2.5-coder:7b) and avoid upgrading silently.",
    "Freeze the vector store build; do not rebuild between comparison runs.",
    "Record model version, seed, and config hash in every report for auditability.",
]
for b in rec_bullets:
    doc.add_paragraph(b, style="List Bullet")

doc.add_paragraph(
    "Caveat: Exact byte-level reproducibility also depends on Ollama version, "
    "hardware (CPU vs GPU), and numerical precision. Day-to-day runs on the same machine "
    "with the same model tag should be effectively identical after applying the settings above."
)


# Save document
doc.save(output_path)
print(f"Created: {output_path}")
